import streamlit as st
import os
import uuid
import re
import sys
import pathlib

# ✅ Resolve project root dynamically
CURRENT_FILE = pathlib.Path(__file__).resolve()
PROJECT_ROOT = CURRENT_FILE.parent.parent
SRC_PATH = PROJECT_ROOT / "src"
if str(SRC_PATH) not in sys.path:
    sys.path.insert(0, str(SRC_PATH))

from Helpers.utility import check_password_multi

st.title("MDWcare Assistant — Document Upload")

# ✅ Check login
if not st.session_state.get("password_correct"):
    st.error("🚫 You must log in first.")
    st.stop()
else:
    st.success(f"👋 Welcome, **{st.session_state['username']}**!")

if "uploaded_docs" not in st.session_state:
    st.session_state["uploaded_docs"] = []

# --- Helper Functions ---
def clean_for_mvp(text):
    text = re.sub(r"\d{1,2}/\d{1,2}/\d{2,4},? ?\d{1,2}:\d{2} ?[APM]{2}", " ", text)
    text = re.sub(r"Page \d+ of \d+", " ", text)
    text = re.sub(r"[^\w\s.,?!:;/=&%\-]", "", text)
    text = re.sub(r"\s+", " ", text)
    lines = [line.strip() for line in text.split("\n") if len(line.strip().split()) > 3]
    return "\n".join(lines)

def chunk_text(text, chunk_size):
    return [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]

def extract_text_from_file(file, filetype):
    text = ""
    try:
        if filetype == "pdf":
            try:
                import pdfplumber
                with pdfplumber.open(file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text() or ""
                        page_text = re.sub(r"[^\x00-\x7F]+", " ", page_text)
                        text += page_text + "\n"
            except ImportError:
                st.error("❌ Missing dependency: `pdfplumber`. Run `pip install pdfplumber`.")
        elif filetype == "docx":
            try:
                import docx
                doc = docx.Document(file)
                for para in doc.paragraphs:
                    text += para.text + "\n"
                    for run in para.runs:
                        if "HYPERLINK" in run._element.xml:
                            matches = re.findall(r'https://www\\.mom\\.gov\\.sg[^\s"]+', run._element.xml)
                            for m in matches:
                                text += f"\n{m}\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + "\n"
            except ImportError:
                st.error("❌ Missing dependency: `python-docx`. Run `pip install python-docx`.")
        elif filetype == "txt":
            text = file.read().decode("utf-8")
    except Exception as e:
        st.error(f"❌ Failed to extract text from file: {e}")
    return text

# --- Upload Logic ---
mchunksize = st.slider("Select chunk size (characters)", 100, 1000, 300, step=50)

if st.button("🗑️ Reset uploaded files"):
    st.session_state["uploaded_docs"] = []
    st.success("Session memory for uploaded files has been cleared. You may re-upload now.")

uploaded_files = st.file_uploader("Upload up to 3 documents (PDF, Word, TXT)", type=["pdf", "docx", "txt"], accept_multiple_files=True)

if uploaded_files:
    if len(uploaded_files) > 3:
        st.error("⚠️ You can only upload a maximum of 3 files at once.")
    else:
        for uploaded_file in uploaded_files:
            filename = uploaded_file.name
            if any(doc["filename"] == filename for doc in st.session_state["uploaded_docs"]):
                st.warning(f"⏩ File '{filename}' is already uploaded.")
                continue

            ext = filename.lower().split(".")[-1]
            filetype = ext if ext in ["pdf", "docx", "txt"] else None

            if not filetype:
                st.error(f"❌ Unsupported file type for: {filename}")
                continue

            os.makedirs("data/uploaded_files", exist_ok=True)
            save_path = os.path.join("data", "uploaded_files", filename)
            with open(save_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            st.success(f"📁 File '{filename}' saved.")

            with st.spinner(f"🔍 Extracting text from '{filename}'..."):
                raw_text = extract_text_from_file(uploaded_file, filetype)

            if not raw_text.strip():
                st.warning(f"⚠️ No extractable text found in '{filename}'")
                continue

            cleaned = clean_for_mvp(raw_text)
            chunks = chunk_text(cleaned, mchunksize)
            if not chunks or all(not c.strip() for c in chunks):
                st.error(f"❌ No valid chunks from '{filename}'")
                continue

            file_entry = {
                "filename": filename,
                "chunks": [
                    {
                        "chunk_id": str(uuid.uuid4()),
                        "text": chunk,
                        "source": filename
                    } for chunk in chunks
                ],
                "source": "upload"
            }

            st.session_state["uploaded_docs"].append(file_entry)
            st.success(f"✅ '{filename}' processed with {len(chunks)} chunks.")

        if st.button("Go to Q&A Page ➡️"):
            st.switch_page("pages/2_QA.py")

        for uploaded_file in uploaded_files:
            filename = uploaded_file.name
            st.subheader("📝 Raw Extract Preview")
            st.code(raw_text[:1500] + ("..." if len(raw_text) > 1500 else ""))

        for file_entry in st.session_state["uploaded_docs"][-len(uploaded_files):]:
            st.subheader(f"🧩 All Chunks from `{file_entry['filename']}`")
            for chunk in file_entry["chunks"]:
                st.markdown(f"**{chunk['chunk_id']}**")
                st.code(chunk["text"][:1000] + ("..." if len(chunk["text"]) > 1000 else ""))

# Show uploaded files
if st.session_state["uploaded_docs"]:
    st.write(f"📚 Uploaded documents this session: ({len(st.session_state['uploaded_docs'])})")
    for doc in st.session_state["uploaded_docs"]:
        st.write(f"• {doc['filename']} ({len(doc['chunks'])} chunks)")

st.sidebar.info("Navigate between steps using the sidebar.")

with st.expander("IMPORTANT NOTICE", expanded=True):
    st.markdown("""
    **This web application is a prototype developed for _educational purposes only_.**  
    The information provided here is **NOT intended for real-world usage** and should not be relied upon for making any decisions.

    Furthermore, please be aware that the LLM may generate inaccurate or incorrect information. You assume full responsibility for how you use any generated output.

    Always consult with qualified professionals for accurate and personalized advice.
    """)
